"""Genera COMO_CORRER_LOCAL.docx con instrucciones actualizadas (mayo 2026)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
    p.paragraph_format.left_indent = Cm(0.8)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    p._p.get_or_add_pPr().append(shd)
    return p

def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p

def nota(text):
    p = doc.add_paragraph()
    run = p.add_run('Nota: ' + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
    return p

def step(n, text):
    para(f'  {n}. {text}')

# ── PORTADA ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nComo correr el proyecto en tu PC\n')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Intercambios Academicos - Universidad Piloto de Colombia\n\n')
r2.font.size = Pt(13)

doc.add_paragraph()

# ── 1. REQUISITOS PREVIOS ─────────────────────────────────────────────────────
h('1. Requisitos previos', 1)
para('Antes de comenzar asegurate de tener instalado:')
bullet('Python 3.10 o superior  (descarga en python.org)')
bullet('pip  (viene incluido con Python)')
bullet('Un navegador web moderno (Chrome, Firefox o Edge)')
bullet('Cuenta gratuita en Neon (neon.tech) para la base de datos')
bullet('Cuenta de Gmail con verificacion en dos pasos activada (para enviar codigos 2FA)')
doc.add_paragraph()
para('Para verificar que Python esta instalado, abre una terminal y escribe:')
code('python --version')
para('Debe mostrar Python 3.10 o superior.')

# ── 2. DESCARGAR EL PROYECTO ──────────────────────────────────────────────────
h('2. Descargar el proyecto', 1)
para('Opcion A - Con Git (recomendado):')
code('git clone https://github.com/Juan-Vanegas5/intercambios-academicos.git')
para('Opcion B - Sin Git:')
bullet('Ve a github.com/Juan-Vanegas5/intercambios-academicos')
bullet('Haz clic en el boton verde "Code" -> "Download ZIP"')
bullet('Descomprime el ZIP en una carpeta de tu preferencia')

# ── 3. CONFIGURAR NEON ────────────────────────────────────────────────────────
h('3. Configurar la base de datos en Neon', 1)
para('3.1  Crear el proyecto en Neon', bold=True)
step(1, 'Ve a https://neon.tech e inicia sesion (puedes usar tu cuenta de Google)')
step(2, 'Haz clic en "New Project" y dale cualquier nombre (ej: intercambios-upc)')
step(3, 'Cuando se cree, haz clic en "Connection string" y copia la URL. Luce asi:')
code('postgresql://neondb_owner:clave@ep-algo.us-east-2.aws.neon.tech/neondb?sslmode=require')

doc.add_paragraph()
para('3.2  Crear las tablas (schema.sql)', bold=True)
step(1, 'En el panel de Neon, haz clic en "SQL Editor" en el menu lateral')
step(2, 'Abre el archivo database/schema.sql del proyecto con el Bloc de Notas')
step(3, 'Copia TODO su contenido y pegalo en el SQL Editor de Neon')
step(4, 'Haz clic en "Run" (boton verde). Deben crearse 7 tablas sin errores')

doc.add_paragraph()
para('3.3  Cargar los datos de prueba (seed.sql)', bold=True)
step(1, 'Abre el archivo database/seed.sql del proyecto con el Bloc de Notas')
step(2, 'Copia TODO su contenido y pegalo en el SQL Editor de Neon')
step(3, 'Haz clic en "Run". Deben insertarse convocatorias y usuarios de prueba')
step(4, 'Para verificar: escribe  SELECT * FROM usuarios;  y ejecuta')

# ── 4. CREAR .ENV ─────────────────────────────────────────────────────────────
h('4. Crear el archivo .env del backend', 1)
para('Dentro de la carpeta backend/, crea un archivo nuevo llamado exactamente  .env  (con el punto al inicio) y pega esto:')
code(
    'DATABASE_URL=postgresql://tu_usuario:tu_clave@tu_host/tu_db\n'
    'SECRET_KEY=cualquier_texto_largo_y_secreto_aqui_123\n'
    'EMAIL_USER=tu_correo@gmail.com\n'
    'EMAIL_PASSWORD=contraseña_de_aplicacion_de_16_chars'
)
bullet('DATABASE_URL: reemplaza con la Connection string que copiaste de Neon')
bullet('SECRET_KEY: puede ser cualquier texto largo (no lo compartas)')
bullet('EMAIL_USER y EMAIL_PASSWORD: ver seccion 8 para obtener estos valores')
nota('Este archivo no se sube a GitHub. Cada integrante del equipo debe crearlo en su PC.')

# ── 5. INSTALAR DEPENDENCIAS ──────────────────────────────────────────────────
h('5. Instalar dependencias de Python', 1)
para('Abre una terminal (CMD o PowerShell) y navega hasta la carpeta backend/:')
code('cd intercambios-academicos\\backend')
para('Luego instala las dependencias:')
code('pip install -r requirements.txt')
para('Esto instala FastAPI, SQLAlchemy, bcrypt, python-dotenv y todo lo necesario. Tarda 1-2 minutos la primera vez.')

# ── 6. INICIAR EL SERVIDOR ────────────────────────────────────────────────────
h('6. Iniciar el servidor backend', 1)
para('Desde la carpeta backend/, ejecuta:')
code('python -m uvicorn main:app --reload --port 8001')
para('Si todo esta bien, veras en la terminal:')
code(
    'INFO:     Uvicorn running on http://127.0.0.1:8001 (Press CTRL+C to quit)\n'
    'INFO:     Started reloader process [xxxx] using WatchFiles'
)
nota('El comando DEBE ejecutarse desde dentro de la carpeta backend/. Si lo ejecutas desde la raiz del proyecto, dara error "Could not import module main".')
doc.add_paragraph()
para('Deja esta terminal abierta. El servidor debe estar corriendo para que el frontend funcione.')

# ── 7. ABRIR EL FRONTEND ──────────────────────────────────────────────────────
h('7. Abrir el frontend en el navegador', 1)
para('El frontend son paginas HTML puras, no necesitan servidor propio. Tienes dos opciones:')
doc.add_paragraph()
para('Opcion A - Live Server en VS Code (recomendado):', bold=True)
step(1, 'Abre VS Code y carga la carpeta del proyecto')
step(2, 'Instala la extension "Live Server" (busca en el panel de extensiones)')
step(3, 'Haz clic derecho sobre frontend/pages/index.html')
step(4, 'Selecciona "Open with Live Server"')
step(5, 'El navegador abrira en http://127.0.0.1:5500/frontend/pages/index.html')
doc.add_paragraph()
para('Opcion B - Abrir el archivo directamente:', bold=True)
para('Navega hasta la carpeta frontend/pages/ y haz doble clic en index.html. Se abre en el navegador.')

# ── 8. CORREO 2FA ─────────────────────────────────────────────────────────────
h('8. Configurar el correo para los codigos 2FA', 1)
para(
    'El sistema envia un codigo de 6 digitos al correo cada vez que alguien inicia sesion. '
    'Para que funcione con tu correo de Gmail, necesitas una "Contrasena de aplicacion":'
)
step(1, 'Inicia sesion en Gmail y ve a: Cuenta de Google -> Seguridad')
step(2, 'Activa la "Verificacion en dos pasos" si no la tienes (es obligatorio)')
step(3, 'En el buscador de configuracion escribe "Contraseñas de aplicaciones"')
step(4, 'Selecciona app: Correo, dispositivo: Otro (escribe "intercambios")')
step(5, 'Haz clic en Generar. Google mostrara una clave de 16 letras (ej: yhxm wpsk pwpu ztqf)')
step(6, 'Copia esa clave (sin espacios) y ponla como EMAIL_PASSWORD en el archivo .env')
doc.add_paragraph()
para('Si no configuras el correo, el codigo igual aparece en la terminal del servidor:', bold=True)
code('CODIGO para admin@upc.edu.co: 847291')
para('Copia ese numero de la terminal y usalo en el formulario de login.')

# ── 9. VERIFICACION FINAL ─────────────────────────────────────────────────────
h('9. Verificacion: todo funcionando', 1)
headers = ['Que probar', 'Donde ir', 'Que debe pasar']
rows_data = [
    ['Swagger (docs de la API)', 'http://localhost:8001/docs', 'Muestra todos los endpoints'],
    ['Lista de convocatorias', 'Abre index.html', '4 convocatorias activas visibles'],
    ['Login de administrador', 'login.html con\nadmin@upc.edu.co / admin', 'Pide codigo 2FA y luego entra al panel'],
    ['Login de estudiante', 'login.html con\njuan.vanegas@upc.edu.co / 1234', 'Redirige al inicio con nombre visible'],
    ['Postularse', 'Login como estudiante -> Postularme', 'Postulacion guardada y visible en Mis postulaciones'],
]
t = doc.add_table(rows=1+len(rows_data), cols=3)
t.style = 'Table Grid'
hrow = t.rows[0]
for i, h_text in enumerate(headers):
    hrow.cells[i].text = h_text
    hrow.cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hrow.cells[i], '1a3a6b')
    hrow.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for ri, rd in enumerate(rows_data):
    r = t.rows[ri+1]
    bg = 'F8FAFF' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate(rd):
        r.cells[ci].text = val
        set_cell_bg(r.cells[ci], bg)
for row in t.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(6)

# ── 10. ERRORES COMUNES ───────────────────────────────────────────────────────
doc.add_paragraph()
h('10. Errores comunes y soluciones', 1)
err_headers = ['Error', 'Solucion']
err_rows = [
    ['"Could not import module main"', 'Estas ejecutando uvicorn desde la carpeta raiz. Entra primero a la carpeta backend/ con:  cd backend'],
    ['"Connection refused" en el frontend', 'El servidor backend no esta corriendo. Ejecuta uvicorn y dejalo abierto en otra terminal'],
    ['"No module named fastapi"', 'No instalaste las dependencias. Ejecuta:  pip install -r requirements.txt  desde backend/'],
    ['Codigo 2FA no llega al correo', 'Busca el codigo en la terminal donde corre uvicorn. Dice: CODIGO para tu@email.com: 123456'],
    ['Error de base de datos', 'Verifica que DATABASE_URL en el .env sea exactamente la Connection string de Neon (sin espacios ni comillas extra)'],
    ['Tablas no encontradas en Neon', 'Ejecuta schema.sql y seed.sql en el SQL Editor de Neon (ver paso 3)'],
]
t2 = doc.add_table(rows=1+len(err_rows), cols=2)
t2.style = 'Table Grid'
hrow2 = t2.rows[0]
for i, h_text in enumerate(err_headers):
    hrow2.cells[i].text = h_text
    hrow2.cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hrow2.cells[i], '1a3a6b')
    hrow2.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for ri, rd in enumerate(err_rows):
    r = t2.rows[ri+1]
    bg = 'F8FAFF' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate(rd):
        r.cells[ci].text = val
        set_cell_bg(r.cells[ci], bg)
for row in t2.rows:
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(10.5)

doc.save('docs/COMO_CORRER_LOCAL.docx')
print('OK: docs/COMO_CORRER_LOCAL.docx generado')
