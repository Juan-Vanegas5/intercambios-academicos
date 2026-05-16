"""Genera documentacion_completa.docx con todo el proyecto actualizado (mayo 2026)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos base ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def heading(text, level=1, color="1a3a6b"):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, '1a3a6b')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        r = t.rows[ri+1]
        bg = 'F8FAFF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            r.cells[ci].text = val
            set_cell_bg(r.cells[ci], bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nSistema de Intercambios Academicos\nUniversidad Piloto de Colombia\n')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Documentacion Tecnica Completa del Proyecto\n')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x37, 0x51, 0x8c)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f'Fecha: {datetime.date.today().strftime("%d de %B de %Y")}\n\n')
r3.font.size = Pt(11)
r3.italic = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Desarrollado por: Juan Andres Vanegas\nPrograma: Ingenieria de Sistemas - Desarrollo Web\nPeriodo: 2026-1')
r4.font.size = Pt(11)
r4.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS
# ═══════════════════════════════════════════════════════════════════════════════
heading('Tabla de Contenidos', 1)
toc = [
    '1. Descripcion General del Proyecto',
    '2. Tecnologias Utilizadas',
    '3. Estructura del Proyecto',
    '4. Base de Datos (PostgreSQL - Neon Cloud)',
    '5. Backend - Python/FastAPI',
    '6. Frontend - HTML/CSS/JavaScript',
    '7. Seguridad: JWT, BCrypt y Autenticacion en Dos Pasos (2FA)',
    '8. API REST - Referencia Completa de Endpoints',
    '9. Swagger UI - Documentacion Interactiva',
    '10. Flujos de Usuario Completos',
    '11. Panel de Administrador',
    '12. Como Ejecutar el Proyecto Localmente',
    '13. Configuracion de Correo para 2FA (Gmail)',
    '14. Credenciales de Prueba',
    '15. Glosario Tecnico',
]
for item in toc:
    bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCION GENERAL
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. Descripcion General del Proyecto', 1)
para(
    'El Sistema de Intercambios Academicos es una aplicacion web full-stack desarrollada '
    'para la Universidad Piloto de Colombia (UPC). Permite a los estudiantes consultar '
    'convocatorias de intercambio internacional, postularse a ellas y hacer seguimiento '
    'del estado de sus solicitudes. Los administradores pueden gestionar y aprobar o '
    'rechazar postulaciones desde un panel de control centralizado.'
)

heading('Objetivo del Sistema', 2)
for item in [
    'Digitalizar y centralizar el proceso de intercambios academicos de la UPC',
    'Dar visibilidad a estudiantes sobre oportunidades disponibles',
    'Permitir al equipo administrativo gestionar postulaciones eficientemente',
    'Registrar el historial de decisiones con comentarios de retroalimentacion',
    'Garantizar la seguridad de datos con contraseñas hasheadas y autenticacion 2FA',
]:
    bullet(item)

heading('Caracteristicas Principales', 2)
for item in [
    'Autenticacion segura en dos pasos: contrasena + codigo de 6 digitos enviado al correo',
    'Contraseñas almacenadas con BCrypt (hash irreversible)',
    'Roles diferenciados: Estudiante y Administrador',
    'Gestion de convocatorias por estado: Activa, Proximamente, Cerrada',
    'Panel administrativo con estadisticas en tiempo real',
    'Seguimiento del estado de cada postulacion con comentarios del admin',
    'Subida de documentos PDF (certificado de notas, paz y salvo academico)',
    'Visualizacion y descarga de documentos desde el panel de administrador',
    'Registro de estudiantes con datos completos: cedula, celular, programa academico',
    'Formulario de postulacion con semestre actual y carta de intencion',
    'Documentacion automatica de la API con Swagger UI',
]:
    bullet(item)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. TECNOLOGIAS
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. Tecnologias Utilizadas', 1)

heading('Backend', 2)
simple_table(
    ['Tecnologia', 'Version', 'Proposito'],
    [
        ['Python', '3.14+', 'Lenguaje de programacion principal'],
        ['FastAPI', '0.115+', 'Framework web para la API REST'],
        ['SQLAlchemy', '2.x', 'ORM para mapear tablas a clases Python'],
        ['Pydantic', '2.x', 'Validacion de datos y esquemas de entrada/salida'],
        ['PyJWT', '2.x', 'Generacion y validacion de tokens JWT'],
        ['bcrypt', '4.x', 'Hash de contraseñas (algoritmo BCrypt)'],
        ['python-multipart', 'latest', 'Recepcion de archivos PDF subidos'],
        ['uvicorn', 'latest', 'Servidor ASGI para correr FastAPI'],
        ['python-dotenv', 'latest', 'Carga de variables de entorno desde .env'],
    ],
    [5, 3, 8]
)

heading('Frontend', 2)
simple_table(
    ['Tecnologia', 'Proposito'],
    [
        ['HTML5', 'Estructura de las paginas web'],
        ['CSS3 personalizado', 'Estilos, layout responsive, componentes UI'],
        ['JavaScript (vanilla)', 'Logica del cliente, llamadas a la API, manejo de sesion'],
        ['Fetch API', 'Peticiones HTTP asincronas al backend'],
        ['localStorage', 'Almacenamiento del token JWT y datos del usuario en el navegador'],
    ],
    [5, 11]
)

heading('Base de Datos e Infraestructura', 2)
simple_table(
    ['Tecnologia', 'Proposito'],
    [
        ['PostgreSQL 16', 'Motor de base de datos relacional'],
        ['Neon Cloud', 'Hosting serverless de PostgreSQL (gratuito)'],
        ['SQLAlchemy + psycopg2', 'Conexion Python -> PostgreSQL'],
    ],
    [5, 11]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. ESTRUCTURA DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. Estructura del Proyecto', 1)
para('El proyecto sigue una arquitectura separada: backend (API REST) y frontend (paginas HTML estaticas), con la base de datos alojada en Neon Cloud.')
doc.add_paragraph()
code_block(
    'intercambios-academicos/\n'
    '|-- backend/                   <- API REST en Python/FastAPI\n'
    '|   |-- main.py                <- Punto de entrada, login y registro\n'
    '|   |-- database.py            <- Conexion a PostgreSQL (Neon)\n'
    '|   |-- models.py              <- Clases ORM (tablas de la BD)\n'
    '|   |-- schemas.py             <- Esquemas Pydantic (validacion)\n'
    '|   |-- auth.py                <- BCrypt y generacion de JWT\n'
    '|   |-- email_service.py       <- Envio de codigo 2FA por correo\n'
    '|   |-- .env                   <- Credenciales (NO subir a Git)\n'
    '|   |-- requirements.txt       <- Dependencias Python\n'
    '|   |-- uploads/               <- Archivos PDF subidos por estudiantes\n'
    '|   `-- routers/\n'
    '|       |-- convocatorias.py   <- GET /api/convocatorias\n'
    '|       |-- postulaciones.py   <- POST/GET /api/postulaciones\n'
    '|       `-- admin.py           <- Endpoints del administrador\n'
    '|\n'
    '|-- frontend/\n'
    '|   |-- pages/\n'
    '|   |   |-- index.html         <- Lista de convocatorias (publica)\n'
    '|   |   |-- login.html         <- Login en 2 pasos\n'
    '|   |   |-- registro.html      <- Registro de estudiantes\n'
    '|   |   |-- postulacion.html   <- Formulario de postulacion\n'
    '|   |   |-- seguimiento.html   <- Mis postulaciones\n'
    '|   |   `-- admin/panel.html   <- Panel de administrador\n'
    '|   |-- css/styles.css         <- Estilos globales\n'
    '|   `-- js/main.js             <- Utilidades JS compartidas\n'
    '|\n'
    '|-- database/\n'
    '|   |-- schema.sql             <- Crea todas las tablas\n'
    '|   `-- seed.sql               <- Datos de prueba iniciales\n'
    '|\n'
    '`-- docs/\n'
    '    |-- documentacion_completa.docx\n'
    '    `-- generar_doc.py\n'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. Base de Datos (PostgreSQL - Neon Cloud)', 1)
para(
    'La base de datos esta alojada en Neon (neon.tech), una plataforma serverless de '
    'PostgreSQL gratuita. El proyecto incluye dos archivos SQL para inicializarla:'
)
bullet('database/schema.sql — Crea todas las tablas con sus columnas y relaciones')
bullet('database/seed.sql — Inserta datos de prueba (convocatorias, usuarios, etc.)')

heading('Tablas del Sistema', 2)
simple_table(
    ['Tabla', 'Descripcion', 'Columnas principales'],
    [
        ['programas_academicos', 'Carreras universitarias', 'id, nombre'],
        ['universidades', 'Instituciones de destino', 'id, nombre, pais'],
        ['tipos_documentos', 'Tipos de doc. requeridos', 'id, nombre'],
        ['usuarios', 'Estudiantes y administradores', 'id, nombre, apellido, email, contrasena (BCrypt), rol, codigo, cedula, celular, programa_id'],
        ['convocatorias', 'Programas de intercambio', 'id, titulo, universidad_id, descripcion, requisitos, fecha_inicio, fecha_cierre, cupos, estado'],
        ['postulaciones', 'Solicitudes de estudiantes', 'id, estudiante_id, convocatoria_id, semestre, carta_intencion, estado, comentario_admin, fecha_postulacion'],
        ['documentos', 'Archivos PDF subidos', 'id, postulacion_id, nombre_archivo, tipo_documento_id, ruta_archivo, fecha_subida'],
    ],
    [4, 5, 7]
)

heading('Diagrama de Relaciones (simplificado)', 2)
code_block(
    'usuarios (programa_id) -------> programas_academicos\n'
    'convocatorias (universidad_id) -> universidades\n'
    'postulaciones (estudiante_id) --> usuarios\n'
    'postulaciones (convocatoria_id)-> convocatorias\n'
    'documentos (postulacion_id) ---> postulaciones\n'
    'documentos (tipo_documento_id) -> tipos_documentos\n'
)

heading('Como configurar la BD en Neon', 2)
para('Sigue estos pasos para inicializar la base de datos en Neon:', bold=True)
for i, step in enumerate([
    'Inicia sesion en https://neon.tech y abre tu proyecto',
    'En el menu lateral, haz clic en "SQL Editor"',
    'Copia todo el contenido de database/schema.sql y pegalo en el editor',
    'Haz clic en "Run" (boton verde). Deben crearse 7 tablas sin errores',
    'Luego copia todo el contenido de database/seed.sql y ejecutalo del mismo modo',
    'Para verificar: escribe SELECT * FROM usuarios; y ejecuta — deben aparecer los usuarios de prueba',
], 1):
    para(f'{i}. {step}')

para('IMPORTANTE: Si necesitas reiniciar los datos, seed.sql comienza con TRUNCATE para limpiar todo antes de insertar.', italic=True, color='dc2626')

# ═══════════════════════════════════════════════════════════════════════════════
# 5. BACKEND
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. Backend - Python/FastAPI', 1)
para(
    'El backend expone una API REST que el frontend consume mediante fetch(). '
    'Usa FastAPI como framework principal, SQLAlchemy como ORM y JWT para autenticacion.'
)

heading('Archivo .env (variables de entorno)', 2)
para('El backend lee credenciales desde backend/.env. Este archivo NO se sube a Git por seguridad.')
code_block(
    '# backend/.env\n'
    'DATABASE_URL=postgresql://usuario:contrasena@host/nombre_db\n'
    'SECRET_KEY=clave_secreta_para_jwt\n'
    'EMAIL_USER=tu_correo@gmail.com\n'
    'EMAIL_PASSWORD=contraseña_de_aplicacion_gmail\n'
)
para('DATABASE_URL se obtiene desde el panel de Neon -> Connection string.', italic=True)

heading('Modulo de Autenticacion (auth.py)', 2)
para('Contiene dos funciones criticas:')
code_block(
    'hashear_contrasena(raw: str) -> str\n'
    '    # Genera un hash BCrypt irreversible de la contrasena\n'
    '    # Ejemplo: "1234" -> "$2b$12$jWMY94lZT65qFVcV..."\n\n'
    'verificar_contrasena(raw: str, guardada: str) -> bool\n'
    '    # Compara la contrasena en texto plano con el hash de la BD\n'
    '    # Devuelve True si coincide, False si no\n'
)

heading('Modulo de Email 2FA (email_service.py)', 2)
para('Gestiona los codigos de verificacion de 6 digitos para el login:')
bullet('Los codigos se guardan en memoria con expiracion de 10 minutos')
bullet('Si el correo esta configurado en .env, envia el codigo por Gmail SMTP')
bullet('Si el envio falla (correo incorrecto, sin internet), imprime el codigo en la terminal del servidor como respaldo')
bullet('El proceso de login siempre continua, nunca bloquea por fallo de correo')

heading('Routers (endpoints por modulo)', 2)
simple_table(
    ['Archivo', 'Prefijo', 'Responsabilidad'],
    [
        ['main.py', '/api/auth', 'Login 2 pasos, registro, verificacion de codigo'],
        ['routers/convocatorias.py', '/api/convocatorias', 'Listar convocatorias activas (publico)'],
        ['routers/postulaciones.py', '/api/postulaciones', 'Crear postulacion, subir docs, ver las propias'],
        ['routers/admin.py', '/api/admin', 'Ver todas las postulaciones, cambiar estado, ver documentos'],
    ],
    [5, 5, 6]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. FRONTEND
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. Frontend - HTML/CSS/JavaScript', 1)
para(
    'El frontend es un conjunto de paginas HTML estaticas que se comunican con el backend '
    'mediante la Fetch API de JavaScript. No usa ningun framework frontend (sin React, Vue, etc.).'
)

heading('Paginas del Sistema', 2)
simple_table(
    ['Pagina', 'Acceso', 'Descripcion'],
    [
        ['index.html', 'Publica', 'Lista convocatorias activas. Si el usuario esta logueado muestra botones de postulacion; si no, muestra login/registro'],
        ['login.html', 'Publica', 'Login en dos pasos: primero email+contrasena, luego codigo de 6 digitos enviado al correo'],
        ['registro.html', 'Publica', 'Formulario de registro con nombre, apellido, email, cedula, celular, codigo estudiantil, programa y contrasena'],
        ['postulacion.html', 'Estudiante', 'Formulario de postulacion con selector de convocatoria, semestre actual, carta de intencion y subida de PDFs'],
        ['seguimiento.html', 'Estudiante', 'Lista de postulaciones propias con estado (en revision, aprobada, rechazada) y comentario del admin'],
        ['admin/panel.html', 'Admin', 'Panel con todas las postulaciones, estadisticas y modal de detalle con datos del estudiante y documentos'],
    ],
    [4, 3, 9]
)

heading('main.js - Utilidades Compartidas', 2)
para('El archivo frontend/js/main.js define funciones usadas por todas las paginas:')
simple_table(
    ['Funcion', 'Descripcion'],
    [
        ['BASE_URL', 'URL base del backend: http://localhost:8001'],
        ['apiFetch(url, opts)', 'Wrapper de fetch que añade BASE_URL, maneja errores HTTP y parsea JSON'],
        ['authHeaders()', 'Devuelve { Content-Type, Authorization: Bearer <token> }'],
        ['getToken()', 'Lee el JWT del localStorage'],
        ['getUsuario()', 'Lee los datos del usuario (nombre, rol, etc.) del localStorage'],
        ['setSession(token, usuario)', 'Guarda token y datos de usuario en localStorage'],
        ['cerrarSesion()', 'Borra localStorage y redirige al login'],
        ['requireAuth()', 'Si no hay sesion, redirige al login'],
        ['requireAdmin()', 'Si el usuario no es admin, redirige al inicio'],
        ['formatFecha(str)', 'Formatea una fecha ISO a DD/MM/YYYY'],
        ['actualizarNav()', 'Muestra el nombre del usuario en la barra de navegacion'],
        ['mostrarError(div, msg)', 'Muestra un mensaje de error en un elemento HTML'],
    ],
    [5, 11]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════════
heading('7. Seguridad: JWT, BCrypt y Autenticacion en Dos Pasos (2FA)', 1)

heading('BCrypt - Hash de Contraseñas', 2)
para(
    'Las contraseñas NUNCA se guardan en texto plano en la base de datos. '
    'Al registrarse o crear un usuario, la contrasena pasa por el algoritmo BCrypt '
    'que genera un hash irreversible de 60 caracteres.'
)
code_block(
    'Contraseña original: "1234"\n'
    'Hash BCrypt en BD:   "$2b$12$jWMY94lZT65qFVcVvQgdN.DXcu96JDpgRQy54c7wmivF1AKQ4WqBu"\n\n'
    'Ventajas de BCrypt:\n'
    '- Irreversible: no se puede obtener "1234" a partir del hash\n'
    '- Salt automatico: dos hashes del mismo texto son diferentes\n'
    '- Factor de costo (12): hace el hash lento, dificultando ataques por fuerza bruta\n'
)

heading('JWT - Token de Sesion', 2)
para('Tras el login exitoso, el servidor genera un JWT (JSON Web Token) con:')
bullet('Email del usuario (sub)')
bullet('Rol: "estudiante" o "admin"')
bullet('Expiracion: 24 horas (exp)')
para('El frontend guarda el token en localStorage y lo envia en cada peticion protegida:')
code_block('Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...')
para('El backend valida la firma del token con la SECRET_KEY del .env. Si el token es invalido o expirado, devuelve 401 Unauthorized.')

heading('2FA - Autenticacion en Dos Pasos', 2)
para('El proceso de login tiene dos etapas obligatorias:')
simple_table(
    ['Paso', 'Accion del Usuario', 'Respuesta del Servidor'],
    [
        ['1', 'Ingresa email y contrasena en login.html', 'Verifica credenciales y envia codigo de 6 digitos al correo. Devuelve {"mensaje": "Codigo enviado", "email": "..."}'],
        ['2', 'Ingresa el codigo de 6 digitos recibido', 'Verifica el codigo (valido por 10 min). Si es correcto, devuelve el JWT y datos del usuario.'],
    ],
    [1.5, 6, 8.5]
)
para('Si el correo no esta configurado, el codigo se imprime en la terminal del servidor (util para pruebas locales).', italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. API REST
# ═══════════════════════════════════════════════════════════════════════════════
heading('8. API REST - Referencia Completa de Endpoints', 1)

heading('Autenticacion', 2)
simple_table(
    ['Metodo', 'Endpoint', 'Auth', 'Descripcion', 'Body / Response'],
    [
        ['POST', '/api/auth/login', 'No', 'Paso 1: verificar credenciales y enviar codigo 2FA', 'Body: {email, contrasena}\nResp: {mensaje, email}'],
        ['POST', '/api/auth/verificar-codigo', 'No', 'Paso 2: ingresar codigo y obtener JWT', 'Body: {email, codigo}\nResp: {token, nombre, apellido, email, rol}'],
        ['POST', '/api/auth/registro', 'No', 'Crear cuenta de estudiante', 'Body: {nombre, apellido, email, contrasena, cedula, celular, codigo, programa}\nResp: {token, nombre, email, rol}'],
    ],
    [2, 5, 1.5, 4, 5]
)

heading('Convocatorias (publico)', 2)
simple_table(
    ['Metodo', 'Endpoint', 'Auth', 'Descripcion'],
    [
        ['GET', '/api/convocatorias', 'No', 'Lista todas las convocatorias con estado activa o proximamente'],
    ],
    [2, 5, 1.5, 8]
)

heading('Postulaciones (estudiante autenticado)', 2)
simple_table(
    ['Metodo', 'Endpoint', 'Auth', 'Descripcion'],
    [
        ['POST', '/api/postulaciones', 'JWT', 'Crear postulacion. Body: {convocatoriaId, semestre, cartaIntencion}'],
        ['POST', '/api/postulaciones/{id}/documentos', 'JWT', 'Subir PDFs. Form-data: certificado (PDF), paz_y_salvo (PDF)'],
        ['GET', '/api/postulaciones/mis', 'JWT', 'Ver mis postulaciones con estado y comentarios del admin'],
    ],
    [2, 5, 1.5, 8]
)

heading('Panel de Administrador (rol admin)', 2)
simple_table(
    ['Metodo', 'Endpoint', 'Auth', 'Descripcion'],
    [
        ['GET', '/api/admin/postulaciones', 'JWT Admin', 'Ver todas las postulaciones con datos del estudiante'],
        ['PUT', '/api/admin/postulaciones/{id}/estado', 'JWT Admin', 'Cambiar estado y agregar comentario. Body: {estado, comentario}'],
        ['GET', '/api/admin/postulaciones/{id}/documentos', 'JWT Admin', 'Ver lista de documentos de una postulacion'],
        ['GET', '/api/admin/estadisticas', 'JWT Admin', 'Contadores: total, aprobadas, rechazadas, en revision'],
        ['POST', '/api/admin/convocatorias', 'JWT Admin', 'Crear nueva convocatoria'],
        ['DELETE', '/api/admin/convocatorias/{id}', 'JWT Admin', 'Eliminar una convocatoria'],
    ],
    [2, 5, 2.5, 7]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. SWAGGER UI
# ═══════════════════════════════════════════════════════════════════════════════
heading('9. Swagger UI - Documentacion Interactiva', 1)
para(
    'FastAPI genera automaticamente una interfaz Swagger UI que documenta todos los '
    'endpoints y permite probarlos directamente desde el navegador sin necesidad de Postman.'
)
para('URL de acceso (con el servidor corriendo):', bold=True)
code_block('http://localhost:8001/docs')
para('Tambien disponible en formato ReDoc:')
code_block('http://localhost:8001/redoc')

heading('Como usar Swagger para probar el login', 2)
for i, step in enumerate([
    'Abre http://localhost:8001/docs en el navegador',
    'Busca la seccion "Autenticacion" y haz clic en POST /api/auth/login',
    'Haz clic en "Try it out" y luego edita el body con tu email y contrasena',
    'Haz clic en "Execute" - el servidor enviara el codigo a tu correo',
    'Busca POST /api/auth/verificar-codigo, ingresa el email y el codigo recibido',
    'El servidor devolvera el token JWT. Copialo.',
    'Haz clic en el boton "Authorize" (candado) en la parte superior de Swagger',
    'Pega el token en el campo "Value" con el formato: Bearer eyJ...',
    'Ya puedes probar endpoints protegidos directamente desde Swagger',
], 1):
    para(f'{i}. {step}')

# ═══════════════════════════════════════════════════════════════════════════════
# 10. FLUJOS DE USUARIO
# ═══════════════════════════════════════════════════════════════════════════════
heading('10. Flujos de Usuario Completos', 1)

heading('Flujo: Registro de Estudiante', 2)
code_block(
    '1. Estudiante abre registro.html\n'
    '2. Completa el formulario: nombre, apellido, email institucional,\n'
    '   cedula, celular, codigo estudiantil, programa academico, contrasena\n'
    '3. Al enviar, el JS llama POST /api/auth/registro\n'
    '4. El backend hashea la contrasena con BCrypt y guarda el usuario en BD\n'
    '5. Responde con token JWT y datos del usuario\n'
    '6. El frontend guarda la sesion en localStorage y redirige a index.html\n'
)

heading('Flujo: Login con 2FA', 2)
code_block(
    '1. Estudiante abre login.html - ve formulario de email y contrasena\n'
    '2. Ingresa credenciales y hace clic en "Continuar"\n'
    '3. Frontend llama POST /api/auth/login\n'
    '4. Backend verifica email y contrasena con BCrypt\n'
    '5. Si son correctas, genera codigo de 6 digitos y lo guarda (10 min TTL)\n'
    '6. Envia el codigo al correo del usuario via Gmail SMTP\n'
    '7. Responde: {"mensaje": "Codigo enviado", "email": "..."}\n'
    '8. Frontend muestra el campo para ingresar el codigo (paso 2)\n'
    '9. Usuario ingresa el codigo y hace clic en "Verificar"\n'
    '10. Frontend llama POST /api/auth/verificar-codigo\n'
    '11. Backend valida el codigo y su expiracion\n'
    '12. Si es valido, genera JWT y responde con token y datos del usuario\n'
    '13. Frontend guarda sesion en localStorage y redirige a index.html\n'
)

heading('Flujo: Postulacion a una Convocatoria', 2)
code_block(
    '1. Estudiante ve convocatorias en index.html y hace clic en "Postularme"\n'
    '2. Se redirige a postulacion.html con el ID de la convocatoria en la URL\n'
    '3. El formulario pre-selecciona la convocatoria y muestra sus requisitos\n'
    '4. Estudiante selecciona su semestre actual\n'
    '5. Redacta la carta de intencion (programa destino, motivos, datos personales)\n'
    '6. Adjunta los PDFs requeridos: certificado de notas y paz y salvo academico\n'
    '7. Al enviar:\n'
    '   a. POST /api/postulaciones (crea la postulacion con semestre y carta)\n'
    '   b. POST /api/postulaciones/{id}/documentos (sube los PDFs)\n'
    '8. Exito: redirige a seguimiento.html para ver el estado\n'
)

heading('Flujo: Revision Administrativa', 2)
code_block(
    '1. Admin inicia sesion y llega al panel (admin/panel.html)\n'
    '2. Ve la tabla con todas las postulaciones y estadisticas globales\n'
    '3. Hace clic en el boton de detalle (lupa) de una postulacion\n'
    '4. El modal muestra:\n'
    '   - Datos del estudiante: nombre, programa, cedula, celular, semestre\n'
    '   - Convocatoria a la que aplico\n'
    '   - Carta de intencion completa\n'
    '   - Lista de documentos con links de descarga (PDFs)\n'
    '5. Admin revisa todo y puede:\n'
    '   a. Escribir un comentario de retroalimentacion\n'
    '   b. Hacer clic en "Aprobar" o "Rechazar"\n'
    '6. El sistema llama PUT /api/admin/postulaciones/{id}/estado\n'
    '7. El estudiante puede ver el nuevo estado en seguimiento.html\n'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 11. PANEL DE ADMINISTRADOR
# ═══════════════════════════════════════════════════════════════════════════════
heading('11. Panel de Administrador', 1)
para('El panel de administrador (admin/panel.html) es accesible solo con rol "admin". Incluye:')

heading('Estadisticas en Tiempo Real', 2)
para('Al cargar la pagina, consulta GET /api/admin/estadisticas y muestra:')
bullet('Total de postulaciones recibidas')
bullet('Postulaciones aprobadas')
bullet('Postulaciones rechazadas')
bullet('Postulaciones en revision')

heading('Tabla de Postulaciones', 2)
para('Columnas: #, Estudiante, Programa, Convocatoria, Semestre, Docs (cantidad), Fecha, Estado, Acciones')
bullet('El estado se muestra con color: verde (aprobada), rojo (rechazada), amarillo (en revision)')
bullet('La columna Docs muestra el numero de documentos subidos')
bullet('El boton de lupa abre el modal de detalle')

heading('Modal de Detalle', 2)
bullet('Datos personales del estudiante: nombre completo, programa, cedula, celular')
bullet('Semestre actual y convocatoria a la que aplico')
bullet('Carta de intencion completa en un area de texto de solo lectura')
bullet('Lista de documentos con tipo (certificado, paz y salvo) y enlace de descarga')
bullet('Area de texto para escribir comentario de retroalimentacion')
bullet('Botones: Aprobar (verde), Rechazar (rojo), Guardar comentario (gris)')

heading('Gestion de Convocatorias', 2)
bullet('Formulario para crear nuevas convocatorias (titulo, universidad, fechas, cupos, requisitos)')
bullet('Boton de eliminar por convocatoria (con confirmacion)')

# ═══════════════════════════════════════════════════════════════════════════════
# 12. COMO EJECUTAR EL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
heading('12. Como Ejecutar el Proyecto Localmente', 1)
para('Sigue estos pasos en orden para poner en marcha el proyecto completo:', bold=True)

heading('Requisitos Previos', 2)
simple_table(
    ['Requisito', 'Version', 'Como verificar'],
    [
        ['Python', '3.10+', 'python --version'],
        ['pip', 'Incluido con Python', 'pip --version'],
        ['Cuenta en Neon', 'Gratuita', 'neon.tech (registro gratis)'],
        ['Cuenta de Gmail', 'Con 2FA activado', 'Para el codigo de verificacion'],
        ['Navegador web', 'Chrome / Firefox / Edge', '-'],
    ],
    [4, 4, 8]
)

heading('Paso 1: Clonar o descargar el proyecto', 2)
code_block(
    'git clone https://github.com/Juan-Vanegas5/intercambios-academicos.git\n'
    'cd intercambios-academicos\n'
)

heading('Paso 2: Configurar la base de datos en Neon', 2)
para('2.1 Crear el proyecto en Neon:', bold=True)
for step in [
    'Ve a https://neon.tech y crea una cuenta gratuita',
    'Crea un nuevo proyecto (cualquier nombre)',
    'Anota la "Connection string" que tiene este formato:',
]:
    bullet(step)
code_block('postgresql://usuario:contrasena@ep-algo.us-east-2.aws.neon.tech/neondb?sslmode=require')

para('2.2 Ejecutar el schema:', bold=True)
bullet('En Neon, ve a "SQL Editor"')
bullet('Pega el contenido de database/schema.sql y haz clic en Run')
bullet('Pega el contenido de database/seed.sql y haz clic en Run')

heading('Paso 3: Crear el archivo .env del backend', 2)
para('Crea el archivo backend/.env con este contenido:')
code_block(
    'DATABASE_URL=postgresql://tu_usuario:tu_contrasena@tu_host/tu_db\n'
    'SECRET_KEY=una_clave_secreta_larga_y_aleatoria_aqui\n'
    'EMAIL_USER=tu_correo@gmail.com\n'
    'EMAIL_PASSWORD=contraseña_de_aplicacion_16_caracteres\n'
)
para('Sustituye DATABASE_URL con la Connection string de Neon.', italic=True)
para('Para EMAIL_PASSWORD, ve a la seccion 13 de este documento.', italic=True)

heading('Paso 4: Instalar dependencias Python', 2)
code_block(
    '# Abre una terminal y navega a la carpeta backend\n'
    'cd backend\n\n'
    '# Instalar dependencias\n'
    'pip install -r requirements.txt\n'
)

heading('Paso 5: Iniciar el servidor backend', 2)
code_block(
    '# Desde la carpeta backend/\n'
    'python -m uvicorn main:app --reload --port 8001\n'
)
para('Si todo esta bien, veras en la terminal:')
code_block(
    'INFO:     Uvicorn running on http://127.0.0.1:8001 (Press CTRL+C to quit)\n'
    'INFO:     Started reloader process\n'
)
para('IMPORTANTE: El comando debe ejecutarse desde dentro de la carpeta backend/, no desde la raiz del proyecto.', italic=True, color='dc2626')

heading('Paso 6: Abrir el frontend en el navegador', 2)
para('El frontend es HTML puro, no necesita servidor. Tienes dos opciones:')
para('Opcion A - Abrir directo (puede tener restricciones CORS en algunos navegadores):', bold=True)
code_block('Abre directamente el archivo: frontend/pages/index.html')
para('Opcion B - Usar Live Server en VS Code (recomendado):', bold=True)
bullet('Instala la extension "Live Server" en VS Code')
bullet('Haz clic derecho sobre frontend/pages/index.html')
bullet('Selecciona "Open with Live Server"')
bullet('El navegador abrira en http://127.0.0.1:5500/frontend/pages/index.html')

heading('Verificacion Final', 2)
para('Con el servidor corriendo, verifica que todo funciona:', bold=True)
simple_table(
    ['Que verificar', 'Como', 'Resultado esperado'],
    [
        ['Swagger UI', 'Abre http://localhost:8001/docs', 'Muestra todos los endpoints documentados'],
        ['Convocatorias', 'Abre index.html', 'Lista 4 convocatorias activas'],
        ['Login', 'Usa admin@upc.edu.co / admin', 'Pide codigo, verifica y da acceso'],
        ['Panel admin', 'Login como admin -> panel', 'Muestra estadisticas y tabla de postulaciones'],
        ['Postulacion', 'Login como estudiante y postulate', 'Postulacion creada y visible en seguimiento'],
    ],
    [4, 5, 7]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 13. CONFIGURACION DE CORREO
# ═══════════════════════════════════════════════════════════════════════════════
heading('13. Configuracion de Correo para 2FA (Gmail)', 1)
para(
    'Para que el sistema envie los codigos de verificacion por correo real, necesitas '
    'una "Contrasena de Aplicacion" de Gmail. Esto es diferente a tu contrasena normal.'
)

heading('Pasos para obtener la Contrasena de Aplicacion', 2)
for i, step in enumerate([
    'Inicia sesion en tu cuenta de Gmail',
    'Ve a: Gestion de cuenta Google -> Seguridad',
    'Activa la "Verificacion en dos pasos" si no la tienes (obligatorio)',
    'En la seccion Seguridad, busca "Contraseñas de aplicaciones"',
    'Selecciona la app "Correo" y el dispositivo "Otro (nombre personalizado)"',
    'Escribe un nombre como "intercambios-upc" y haz clic en Generar',
    'Google genera una contraseña de 16 caracteres (ej: yhxm wpsk pwpu ztqf)',
    'Copia esa contraseña y usala como EMAIL_PASSWORD en el archivo .env',
    'Nota: los espacios en la contraseña son opcionales, puedes quitarlos',
], 1):
    para(f'{i}. {step}')

para('Importante: La contraseña de aplicacion solo se muestra una vez. Guardala en un lugar seguro.', italic=True, color='dc2626')

heading('Que pasa si el correo no esta configurado', 2)
para(
    'El sistema funciona aunque el correo no este configurado. En ese caso, '
    'el codigo de verificacion se imprime en la terminal del servidor:'
)
code_block(
    'CODIGO para admin@upc.edu.co: 847291\n'
    '(Advertencia: no se pudo enviar el correo)\n'
)
para('Copia el codigo de la terminal y usalo en el formulario de login.', italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# 14. CREDENCIALES DE PRUEBA
# ═══════════════════════════════════════════════════════════════════════════════
heading('14. Credenciales de Prueba', 1)
para('El seed.sql incluye estos usuarios de prueba con contraseñas hasheadas en BCrypt:')

simple_table(
    ['Rol', 'Email', 'Contrasena', 'Nombre'],
    [
        ['Admin', 'admin@upc.edu.co', 'admin', 'Administrador UPC'],
        ['Estudiante', 'juan.vanegas@upc.edu.co', '1234', 'Juan Vanegas'],
        ['Estudiante', 'maria.garcia@upc.edu.co', 'admin123', 'Maria Garcia'],
    ],
    [3, 6, 4, 5]
)

para('Para el admin: el codigo 2FA se imprime en la terminal del servidor (su email no es real).', italic=True)
para('Para los estudiantes: si tienen un correo real configurado en BD, el codigo llega por email.', italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# 15. GLOSARIO
# ═══════════════════════════════════════════════════════════════════════════════
heading('15. Glosario Tecnico', 1)
simple_table(
    ['Termino', 'Definicion'],
    [
        ['API REST', 'Interfaz de comunicacion entre el frontend y backend que usa HTTP y devuelve datos en formato JSON'],
        ['JWT (JSON Web Token)', 'Cadena de texto firmada que prueba la identidad del usuario. El frontend lo guarda y lo envia en cada peticion protegida'],
        ['BCrypt', 'Algoritmo de hash para contraseñas. Genera una cadena irreversible que el servidor compara sin conocer la contrasena original'],
        ['2FA', 'Autenticacion en dos factores: primero contrasena, luego un codigo temporal. Aumenta la seguridad considerablemente'],
        ['ORM', 'Object-Relational Mapping. SQLAlchemy traduce las clases Python a tablas SQL y viceversa, sin escribir SQL manual'],
        ['Endpoint', 'URL especifica de la API que realiza una accion determinada. Ej: POST /api/auth/login'],
        ['Swagger UI', 'Interfaz web generada automaticamente que documenta y permite probar todos los endpoints de la API'],
        ['Neon Cloud', 'Servicio gratuito de PostgreSQL en la nube. Permite tener una BD real sin instalar PostgreSQL localmente'],
        ['CORS', 'Politica del navegador que controla que dominios pueden hacer peticiones a la API. FastAPI lo configura para permitir el frontend'],
        ['localStorage', 'Almacenamiento del navegador donde el frontend guarda el token JWT y los datos del usuario entre paginas'],
        ['Multipart/form-data', 'Tipo de peticion HTTP que permite subir archivos. Usado para enviar los PDFs de los documentos'],
        ['Hash', 'Resultado de aplicar un algoritmo criptografico a un texto. Irreversible: no se puede recuperar el original a partir del hash'],
        ['Salt', 'Cadena aleatoria añadida antes de hashear. Garantiza que dos contraseñas iguales produzcan hashes diferentes'],
        ['Rol', 'Nivel de acceso del usuario: "estudiante" puede postularse; "admin" puede gestionar postulaciones y convocatorias'],
        ['Semestre', 'Periodo academico del estudiante al momento de postularse. La UPC exige estar en semestre 3 o superior para intercambios'],
        ['Carta de intencion', 'Texto redactado por el estudiante explicando sus motivos para el intercambio, programa destino y datos de contacto'],
    ],
    [5, 11]
)

# Guardar
path = 'docs/documentacion_completa.docx'
doc.save(path)
print(f'Documento generado: {path}')
